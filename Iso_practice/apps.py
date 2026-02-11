import os
import re
import streamlit as st
from dotenv import load_dotenv

from langchain.schema import Document
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chat_models import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document as WordDocument

# ==========================================
# Load API Key
# ==========================================
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    st.error("OPENAI_API_KEY not found in .env file")
    st.stop()

# ==========================================
# Streamlit UI
# ==========================================
st.set_page_config(page_title="ISO Word RAG Assistant", layout="wide")
st.title("📘 ISO Word RAG Assistant")
st.write("Upload an ISO Word document and ask questions.")

uploaded_file = st.file_uploader("Upload ISO Word file (.docx)", type=["docx"])

# ==========================================
# Text Cleaning
# ==========================================
def clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^\w\s\.,;:()\-/]", "", text)
    return text.strip()

# ==========================================
# Load Word File
# ==========================================
def load_and_clean_word(uploaded_file):
    doc = WordDocument(uploaded_file)
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    return [Document(page_content=clean_text(full_text), metadata={"source": "Word"})]

# ==========================================
# Build RAG System
# ==========================================
def build_rag(documents):
    # Split documents into chunks
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    splits = splitter.split_documents(documents)

    # Embeddings and Vectorstore
    embeddings = OpenAIEmbeddings()
    vectorstore = FAISS.from_documents(splits, embeddings)
    retriever = vectorstore.as_retriever(search_kwargs={"k": 4})

    # LLM
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.2)

    # Prompt template
    prompt = ChatPromptTemplate.from_template("""
You are an ISO Quality Management expert.

Answer the question using ONLY the provided context.
If the answer is not found, say:
"The requested information is not available in the provided ISO documentation."

Provide:
- A structured paragraph
- Professional tone
- Friendly and formal style

Context:
{context}

Question:
{input}

Answer:
""")

    # Stuff documents chain
    document_chain = create_stuff_documents_chain(llm, prompt)

    # Retrieval chain
    retrieval_chain = create_retrieval_chain(retriever, document_chain)

    return retrieval_chain

# ==========================================
# Process uploaded file
# ==========================================
if uploaded_file:
    if "rag_chain" not in st.session_state:
        with st.spinner("Processing document and building RAG system..."):
            documents = load_and_clean_word(uploaded_file)
            st.session_state.rag_chain = build_rag(documents)
        st.success("✅ RAG system ready!")

    question = st.text_input("Ask a question about the ISO document:")
    if question:
        with st.spinner("Generating answer..."):
            result = st.session_state.rag_chain.run(question)
        st.markdown("### 📌 Answer")
        st.write(result)
